from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

out=Path(__file__).with_name('PROJECT_DOCUMENTATION.docx')
doc=Document(); sec=doc.sections[0]
sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
styles=doc.styles; styles['Normal'].font.name='Calibri'; styles['Normal'].font.size=Pt(11); styles['Normal'].paragraph_format.space_after=Pt(6); styles['Normal'].paragraph_format.line_spacing=1.25
for n,z,c in [('Heading 1',16,'2E74B5'),('Heading 2',13,'2E74B5'),('Heading 3',12,'1F4D78')]:
 s=styles[n]; s.font.name='Calibri'; s.font.size=Pt(z); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(c)
def shade(cell,fill):
 p=cell._tc.get_or_add_tcPr(); e=OxmlElement('w:shd'); e.set(qn('w:fill'),fill); p.append(e)
def ct(cell,v,b=False,color=None):
 cell.text=''; r=cell.paragraphs[0].add_run(str(v)); r.bold=b
 if color: r.font.color.rgb=RGBColor.from_string(color)
 cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
def tbl(hs,rows):
 t=doc.add_table(rows=1,cols=len(hs)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
 for i,h in enumerate(hs): ct(t.rows[0].cells[i],h,True,'FFFFFF'); shade(t.rows[0].cells[i],'2E74B5')
 for j,row in enumerate(rows):
  cs=t.add_row().cells
  for i,v in enumerate(row): ct(cs[i],v); shade(cs[i],'F4F6F9' if j%2 else 'FFFFFF')
 doc.add_paragraph()
def bl(items):
 for x in items: doc.add_paragraph(x,style='List Bullet')
def num(items):
 for x in items: doc.add_paragraph(x,style='List Number')

h=sec.header.paragraphs[0]; h.text='PHẾ LIỆU THIÊN NHUNG  |  TÀI LIỆU DỰ ÁN'; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT; h.runs[0].font.size=Pt(8)
f=sec.footer.paragraphs[0]; f.text='Tài liệu kỹ thuật nội bộ'; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; f.runs[0].font.size=Pt(8)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Phế Liệu Thiên Nhung'); r.bold=True; r.font.size=Pt(26); r.font.color.rgb=RGBColor.from_string('17352E')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Tài liệu đặc tả, vận hành và kiểm thử dự án'); r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string('E46D3D')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Node.js · React/Vinext · PostgreSQL').font.size=Pt(10)
p=doc.add_paragraph(); p.add_run('Tóm tắt: ').bold=True; p.add_run('Ứng dụng quản lý thu mua phế liệu gồm website công khai, khu vực admin và cổng customer. Backend PostgreSQL xử lý authentication, RBAC, CRUD, đơn hàng, tồn kho, hóa đơn và báo cáo.')

doc.add_heading('1. Phạm vi và mục tiêu',1); doc.add_paragraph('Hệ thống hỗ trợ hoạt động thu mua phế liệu từ bảng giá công khai đến quản lý khách hàng, tạo đơn, cập nhật tồn kho, xuất hóa đơn và theo dõi báo cáo.')
bl(['Admin quản lý khách hàng, vật liệu, giá, đơn hàng, tồn kho, hóa đơn và báo cáo.','Customer đăng ký bằng số điện thoại, đăng nhập và chỉ xem lịch sử giao dịch của chính mình.','Giao dịch chính dùng PostgreSQL transaction, row lock, validation và audit log.'])
doc.add_heading('2. Kiến trúc và cấu trúc source',1)
tbl(['Lớp','Vị trí','Trách nhiệm'],[('Frontend','frontend/ và app/','React/Vinext pages, components, API client, CSS và route entry.'),('Backend','backend/server.mjs','HTTP API, auth/session, RBAC, đơn hàng, tồn kho và báo cáo.'),('Database','PostgreSQL','Users, sessions, customers, materials, prices, orders, inventory, invoices, audit logs.'),('Migration','backend/migrate.mjs','Tạo/cập nhật schema PostgreSQL.'),('Script DBeaver','database.postgres.sql','Schema và dữ liệu mẫu.')])
bl(['backend/ là runtime Node.js/PostgreSQL.','frontend/ là pages, components, API client và styles.','app/ là adapter route bắt buộc của Vinext.','tests/ chứa contract tests và real-DB report tests.'])
doc.add_heading('3. Chức năng nghiệp vụ',1)
for title,items in [('Authentication và RBAC',['Admin đăng nhập email/mật khẩu.','Customer đăng ký bằng số điện thoại.','Session dùng token hash; cookie session HttpOnly, SameSite=Lax.','RBAC và ownership được kiểm tra ở backend.']),('Customer',['Tạo, cập nhật, liên kết user và xóa an toàn.','Customer đã có đơn không bị xóa để bảo toàn lịch sử.']),('Materials và prices',['CRUD vật liệu và bảng giá lịch sử.','Public API hiển thị bảng giá.']),('Orders, inventory và invoices',['Tạo đơn theo customer và vật liệu.','Transaction khóa tồn kho, ghi purchase_items và inventory_movements.','Complete/cancel/update có logic hoàn tác tồn kho.','Invoice lưu snapshot customer và line items.']),('Reports',['Admin: doanh thu, đơn, khối lượng, tồn kho, khách hàng, top vật liệu, ngày/tháng.','Customer: tổng đơn, tiền, kg, vật liệu chính và breakdown ngày/tháng.'])]:
 doc.add_heading(title,2); bl(items)
doc.add_heading('4. API chính',1)
tbl(['Nhóm','Endpoint tiêu biểu','Quyền'],[('Health/public','GET /api/health, /api/materials, /api/prices','Công khai'),('Auth','POST /api/auth/register, /login, /logout; GET /me','Session'),('Customers','GET/POST/PATCH/DELETE /api/customers','Admin'),('Materials/prices','GET public; POST/PATCH/DELETE write','Admin write'),('Orders','GET/POST/PATCH/DELETE /api/orders','Admin'),('Inventory','GET /api/inventory; POST /adjust','Admin'),('Invoices','GET /api/invoices và /:code','Admin'),('Customer','GET /api/customer/orders, /:code, /reports','Customer sở hữu'),('Reports','GET /api/reports/summary, /dashboard','Admin')])
doc.add_heading('5. Bảo mật và toàn vẹn dữ liệu',1)
bl(['SQL tham số hóa.','RBAC và customer ownership ở backend.','CSRF yêu cầu cookie csrf và header x-csrf-token khớp; CORS giới hạn origin.','Input được normalize và giới hạn độ dài.','Transaction rollback trước lỗi và release connection trong finally.','Update/delete kiểm tra affected rows.','Login có rate limit theo IP.','Customer report chỉ tính completed orders của user hiện tại.'])
doc.add_heading('6. Cài đặt và chạy',1); doc.add_heading('Cấu hình .env',2)
q=doc.add_paragraph(); q.paragraph_format.left_indent=Inches(.25); rr=q.add_run('DATABASE_URL=postgresql://postgres:123456@localhost:5432/phe_lieu\nPORT=4000\nNODE_ENV=development\nFRONTEND_ORIGIN=http://localhost:3000'); rr.font.name='Consolas'; rr.font.size=Pt(9)
doc.add_heading('Lệnh',2); num(['npm install','npm run db:migrate','npm run backend','Mở terminal khác và chạy npm run dev.','npm test','npm run test:reports-realdb'])
doc.add_heading('7. Kiểm thử và trạng thái bàn giao',1)
tbl(['Hạng mục','Kết quả','Ghi chú'],[('Build frontend','PASS','Vinext build hoàn tất.'),('Contract/security tests','12/12 PASS','Auth, RBAC, ownership, validation, transaction, reports.'),('Real PostgreSQL reports','3/3 PASS','Admin summary/dashboard và customer scope.'),('Real API health/public','PASS','Health, materials, prices trả dữ liệu thật.'),('Admin/customer E2E','PASS sau fix','Auth, CRUD, order, inventory, invoice và reports.'),('Browser click-through','Cần thủ công','Browser môi trường không mở local navigation.')])
doc.add_paragraph('Các lỗi runtime đã phát hiện và sửa: report SQL trả 500, customer CRUD dùng cột active không tồn tại, customer delete không xử lý quan hệ đơn và payload write chưa đủ chặt.')
doc.add_heading('8. Vận hành và giới hạn',1); bl(['Không commit .env; file đã được gitignore.','Production phải đổi mật khẩu admin mặc định và cấu hình FRONTEND_ORIGIN đúng domain.','Nên dùng database test riêng cho E2E.','Xóa material là soft-delete; xóa customer có đơn bị chặn.','Một số endpoint có giới hạn số bản ghi.','Lint frontend còn một số cảnh báo any/useEffect; không ảnh hưởng build/test.'])
doc.add_heading('9. Hướng phát triển',1); num(['Tách server.mjs thành router/service/repository.','Bổ sung test runtime cho cạnh tranh order và thiếu inventory.','Thêm OTP xác minh số điện thoại.','Thêm phân trang, tìm kiếm và lọc.','Cải thiện responsive/mobile và dọn lint.','Thiết lập CI với database test và backup tự động.'])
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); p.add_run('Trạng thái: ').bold=True; p.add_run('Các flow chính đã được sửa và kiểm thử với PostgreSQL thật; build và test suite đều xanh.')
doc.save(out); print(out)
